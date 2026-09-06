import os
import sys
import glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak, KeepTogether
from reportlab.pdfgen import canvas

def create_exact_pdf(pdf_path, media_files):
    class ExactKECPageCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations()
                super().showPage()
            super().save()

        def draw_page_decorations(self):
            self.saveState()
            
            # Draw Header Logobar Text & Logos representation
            # KEC Green/Blue Header accents
            self.setFont("Helvetica-Bold", 12)
            self.setFillColor(colors.HexColor('#005580'))
            self.drawString(30, A4[1] - 32, "KEC")
            self.setFont("Helvetica-Bold", 8)
            self.setFillColor(colors.HexColor('#666666'))
            self.drawString(30, A4[1] - 40, "TRANSFORM YOURSELF")

            # Center SIH Header
            self.setFont("Helvetica-Bold", 10)
            self.setFillColor(colors.HexColor('#002060'))
            self.drawCentredString(A4[0]/2, A4[1] - 32, "SMART INDIA HACKATHON 2026 @ KONGU ENGINEERING COLLEGE")
            self.setFont("Helvetica-Bold", 9)
            self.setFillColor(colors.HexColor('#111111'))
            self.drawCentredString(A4[0]/2, A4[1] - 44, "(7th & 8th September 2026)")

            # Top Right KONGU Logo text
            self.setFont("Helvetica-Bold", 14)
            self.setFillColor(colors.HexColor('#008080'))
            self.drawRightString(A4[0] - 30, A4[1] - 32, "KONGU")
            self.setFont("Helvetica", 7)
            self.setFillColor(colors.HexColor('#333333'))
            self.drawRightString(A4[0] - 30, A4[1] - 40, "Assuring the best")

            # Divider line below header banner
            self.setStrokeColor(colors.HexColor('#000000'))
            self.setLineWidth(1.0)
            self.line(25, A4[1] - 48, A4[0] - 25, A4[1] - 48)

            # Draw Page Outer Black Frame (exact double black border as in PDF screenshot)
            self.setLineWidth(2.5)
            self.setStrokeColor(colors.HexColor('#000000'))
            self.rect(25, 25, A4[0] - 50, A4[1] - 80)
            self.setLineWidth(1.0)
            self.rect(27, 27, A4[0] - 54, A4[1] - 84)

            self.restoreState()

    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        leftMargin=35,
        rightMargin=35,
        topMargin=58,
        bottomMargin=35
    )

    story = []

    # TOP TABLE (Exact 2-column layout as in template screenshot)
    t_left_html = (
        "<b>Department:</b> Computer Science and Engineering<br/><br/>"
        "<b>Project category:</b> Software<br/><br/>"
        "<b>PS ID:</b> SIH26003<br/><br/>"
        "<b>Team ID:</b> MemoMate"
    )

    t_right_html = (
        "<b>Team leader:</b><br/>"
        "Sriman Kumar V: 24CSR301<br/><br/>"
        "<b>Team members:</b><br/>"
        "Sriman Kumar V: 24CSR301<br/>"
        "Srimathi D: 22CSR180<br/><br/>"
        "<b>Mentor/Co-mentors:</b><br/>"
        "Dr. M. Geetha M.E., Ph.D.: Associate Professor; Department of Computer Science and Engineering, Kongu Engineering College"
    )

    p_tl = Paragraph(t_left_html, ParagraphStyle('TL', fontName='Times-Roman', fontSize=10.5, leading=14))
    p_tr = Paragraph(t_right_html, ParagraphStyle('TR', fontName='Times-Roman', fontSize=10, leading=13))

    tbl = Table([[p_tl, p_tr]], colWidths=[255, 265])
    tbl.setStyle(TableStyle([
        ('BOX', (0, 0), (-1, -1), 2.5, colors.HexColor('#000000')),
        ('INNERGRID', (0, 0), (-1, -1), 2.0, colors.HexColor('#000000')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#FFFFFF'))
    ]))
    story.append(tbl)
    story.append(Spacer(1, 10))

    # Screenshots Insertion Section
    style_cap = ParagraphStyle(
        'CapText',
        fontName='Times-BoldItalic',
        fontSize=10,
        leading=12,
        alignment=1,
        textColor=colors.HexColor('#002060'),
        spaceBefore=3,
        spaceAfter=8
    )

    captions = [
        "Figure 1: MemoMate Personalized User Dashboard & Daily Reminders Interface",
        "Figure 2: Cognitive Metrics & 30-Day Performance Analytics Dashboard",
        "Figure 3: Interactive 3D Mind Matrix Sanctum Module (Three.js WebGL Engine)",
        "Figure 4: Family Memory Recognition & Reminiscence Exercise Module",
        "Figure 5: MemoMate Cognitive Gaming Hub Suite (9 Interactive Games)"
    ]

    for idx, cap in enumerate(captions):
        img_path = media_files[idx] if idx < len(media_files) else None
        if img_path and os.path.exists(img_path):
            img_w = 505
            img_h = 230 if idx in [0, 1, 3] else (290 if idx == 4 else 200)
            
            group = [
                Image(img_path, width=img_w, height=img_h),
                Paragraph(cap, style_cap)
            ]
            story.append(KeepTogether(group))

    doc.build(story, canvasmaker=ExactKECPageCanvas)
    print(f"SUCCESS: Created exact PDF -> {pdf_path}")

def create_exact_docx(docx_path, media_files):
    doc = Document()
    
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    # Header banner
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_hdr.paragraph_format.space_before = Pt(4)
    p_hdr.paragraph_format.space_after = Pt(2)
    r_hdr = p_hdr.add_run("SMART INDIA HACKATHON 2026 @ KONGU ENGINEERING COLLEGE")
    r_hdr.font.name = 'Times New Roman'
    r_hdr.font.size = Pt(13)
    r_hdr.font.bold = True
    r_hdr.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(8)
    r_sub = p_sub.add_run("(7th & 8th September 2026)")
    r_sub.font.name = 'Times New Roman'
    r_sub.font.size = Pt(11)
    r_sub.font.bold = True

    # Top info grid table
    tbl_top = doc.add_table(rows=1, cols=2)
    tbl_top.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_top.autofit = False
    
    cell_left = tbl_top.rows[0].cells[0]
    cell_right = tbl_top.rows[0].cells[1]
    cell_left.width = Inches(3.4)
    cell_right.width = Inches(3.8)

    set_cell_margins(cell_left, top=100, bottom=100, left=120, right=120)
    set_cell_margins(cell_right, top=100, bottom=100, left=120, right=120)
    set_cell_background(cell_left, "FAFAFA")
    set_cell_background(cell_right, "FAFAFA")

    # Left Column
    p_l1 = cell_left.paragraphs[0]
    p_l1.paragraph_format.space_after = Pt(4)
    r = p_l1.add_run("Department: "); r.bold = True; r.font.size = Pt(10.5)
    r2 = p_l1.add_run("Computer Science and Engineering"); r2.font.size = Pt(10.5)

    p_l2 = cell_left.add_paragraph()
    p_l2.paragraph_format.space_after = Pt(4)
    r = p_l2.add_run("Project category: "); r.bold = True; r.font.size = Pt(10.5)
    r2 = p_l2.add_run("Software"); r2.font.size = Pt(10.5)

    p_l3 = cell_left.add_paragraph()
    p_l3.paragraph_format.space_after = Pt(4)
    r = p_l3.add_run("PS ID: "); r.bold = True; r.font.size = Pt(10.5)
    r2 = p_l3.add_run("SIH26003"); r2.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    p_l4 = cell_left.add_paragraph()
    p_l4.paragraph_format.space_after = Pt(4)
    r = p_l4.add_run("Team ID: "); r.bold = True; r.font.size = Pt(10.5)
    r2 = p_l4.add_run("MemoMate"); r2.bold = True; r2.font.size = Pt(10.5)

    # Right Column
    p_r1 = cell_right.paragraphs[0]
    p_r1.paragraph_format.space_after = Pt(2)
    r = p_r1.add_run("Team leader:\n"); r.bold = True; r.font.size = Pt(10.5)
    r2 = p_r1.add_run("Sriman Kumar V: 24CSR301\n\n"); r2.font.size = Pt(10)
    
    r = p_r1.add_run("Team members:\n"); r.bold = True; r.font.size = Pt(10.5)
    m_text = (
        "Sriman Kumar V: 24CSR301\n"
        "Srimathi D: 22CSR180\n\n"
    )
    r2 = p_r1.add_run(m_text); r2.font.size = Pt(9.5)

    r = p_r1.add_run("Mentor/Co-mentors:\n"); r.bold = True; r.font.size = Pt(10.5)
    r2 = p_r1.add_run("Dr. M. Geetha M.E., Ph.D.: Associate Professor; Department of Computer Science and Engineering, Kongu Engineering College"); r2.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    captions = [
        "Figure 1: MemoMate Personalized User Dashboard & Daily Reminders Interface",
        "Figure 2: Cognitive Metrics & 30-Day Performance Analytics Dashboard",
        "Figure 3: Interactive 3D Mind Matrix Sanctum Module (Three.js WebGL Engine)",
        "Figure 4: Family Memory Recognition & Reminiscence Exercise Module",
        "Figure 5: MemoMate Cognitive Gaming Hub Suite (9 Interactive Games)"
    ]

    for idx, cap in enumerate(captions):
        img_path = media_files[idx] if idx < len(media_files) else None
        if img_path and os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(2)
            p_img.paragraph_format.keep_with_next = True
            r_i = p_img.add_run()
            r_i.add_picture(img_path, width=Inches(6.8))

            p_c = doc.add_paragraph()
            p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_c.paragraph_format.space_after = Pt(10)
            p_c.paragraph_format.keep_with_next = True
            r_c = p_c.add_run(cap)
            r_c.font.name = 'Times New Roman'
            r_c.font.size = Pt(10.5)
            r_c.font.bold = True
            r_c.font.italic = True
            r_c.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    doc.save(docx_path)
    print(f"SUCCESS: Created exact Word template -> {docx_path}")

def main():
    media_dir = r'C:\Users\LENOVO\.gemini\antigravity-ide\brain\481a6e73-e62a-4a0d-baee-1a7d30e83efb'
    media_files = [
        os.path.join(media_dir, 'media__1788677112588.png'), # Fig 1
        os.path.join(media_dir, 'media__1788677062655.png'), # Fig 2
        os.path.join(media_dir, 'media__1788677112630.png'), # Fig 3
        os.path.join(media_dir, 'media__1788677112655.png'), # Fig 4
        os.path.join(media_dir, 'media__1788677175164.png')  # Fig 5
    ]

    pdf_path = r'c:\Users\LENOVO\OneDrive\Desktop\SIH\MemoMate.pdf'
    docx_path = r'c:\Users\LENOVO\OneDrive\Desktop\SIH\MemoMate.docx'

    create_exact_pdf(pdf_path, media_files)
    create_exact_docx(docx_path, media_files)

if __name__ == '__main__':
    main()
