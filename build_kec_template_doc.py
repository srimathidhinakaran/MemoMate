import os
import glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak, KeepTogether
from reportlab.pdfgen import canvas

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def build_word_document(docx_path, media_files):
    doc = Document()
    
    # Page setup A4, 0.6 inch margins to match official template border spacing
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    # 1. TEMPLATE HEADER BANNER
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_hdr.paragraph_format.space_before = Pt(4)
    p_hdr.paragraph_format.space_after = Pt(2)
    r_hdr = p_hdr.add_run("SMART INDIA HACKATHON 2026 @ KONGU ENGINEERING COLLEGE")
    r_hdr.font.name = 'Times New Roman'
    r_hdr.font.size = Pt(13)
    r_hdr.font.bold = True
    r_hdr.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(8)
    r_sub = p_sub.add_run("(7th & 8th September 2026)")
    r_sub.font.name = 'Times New Roman'
    r_sub.font.size = Pt(11)
    r_sub.font.bold = True

    # 2. TOP INFORMATION BOX (2-column table)
    tbl_top = doc.add_table(rows=1, cols=2)
    tbl_top.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_top.autofit = False
    
    cell_left = tbl_top.rows[0].cells[0]
    cell_right = tbl_top.rows[0].cells[1]
    cell_left.width = Inches(3.4)
    cell_right.width = Inches(3.8)

    set_cell_margins(cell_left, top=100, bottom=100, left=120, right=120)
    set_cell_margins(cell_right, top=100, bottom=100, left=120, right=120)
    set_cell_background(cell_left, "FAFAFA")
    set_cell_background(cell_right, "FAFAFA")

    # Left Column Content
    p_l1 = cell_left.paragraphs[0]
    p_l1.paragraph_format.space_after = Pt(3)
    r = p_l1.add_run("Department: "); r.bold = True; r.font.size = Pt(10.5)
    r2 = p_l1.add_run("Computer Science and Engineering"); r2.font.size = Pt(10.5)

    p_l2 = cell_left.add_paragraph()
    p_l2.paragraph_format.space_after = Pt(3)
    r = p_l2.add_run("Project category: "); r.bold = True; r.font.size = Pt(10.5)
    r2 = p_l2.add_run("Software"); r2.font.size = Pt(10.5)

    p_l3 = cell_left.add_paragraph()
    p_l3.paragraph_format.space_after = Pt(3)
    r = p_l3.add_run("PS ID: "); r.bold = True; r.font.size = Pt(10.5)
    r2 = p_l3.add_run("SIH26003"); r2.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    p_l4 = cell_left.add_paragraph()
    p_l4.paragraph_format.space_after = Pt(3)
    r = p_l4.add_run("Team ID / Name: "); r.bold = True; r.font.size = Pt(10.5)
    r2 = p_l4.add_run("MemoMate_KEC"); r2.bold = True; r2.font.size = Pt(10.5)

    # Right Column Content
    p_r1 = cell_right.paragraphs[0]
    p_r1.paragraph_format.space_after = Pt(2)
    r = p_r1.add_run("Team leader:\n"); r.bold = True; r.font.size = Pt(10.5)
    r2 = p_r1.add_run("  Srimathi D (Roll No: 22CSR180)\n"); r2.font.size = Pt(10)
    
    r = p_r1.add_run("Team members:\n"); r.bold = True; r.font.size = Pt(10.5)
    m_text = (
        "  Srimathi D (Roll No: 22CSR180)\n"
        "  Team Member 2 (Roll No: 22CSR181)\n"
        "  Team Member 3 (Roll No: 22CSR182)\n"
        "  Team Member 4 (Roll No: 22CSR183)\n"
        "  Team Member 5 (Roll No: 22CSR184)\n"
    )
    r2 = p_r1.add_run(m_text); r2.font.size = Pt(9.5)

    r = p_r1.add_run("Mentor/Co-mentors:\n"); r.bold = True; r.font.size = Pt(10.5)
    r2 = p_r1.add_run("  Dr. R. Malathi; Associate Professor, Dept of CSE, KEC\n  Prof. P. Suresh; Assistant Professor, Dept of CSE, KEC"); r2.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Title section for Prototype Screenshots
    p_sec = doc.add_paragraph()
    p_sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sec.paragraph_format.space_after = Pt(8)
    r_sec = p_sec.add_run("PROTOTYPE SCREENSHOTS & IMPLEMENTATION MODULES")
    r_sec.font.name = 'Times New Roman'
    r_sec.font.size = Pt(13)
    r_sec.font.bold = True
    r_sec.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    # Embed Screenshots with descriptions
    captions = [
        ("Figure 1: MemoMate Personalized User Dashboard & Daily Schedule Interface",
         "Features personalized greeting ('Good Afternoon, Srimathi'), direct baseline assessment launch button, AI voice assistant entry point, accessibility toolbar (English/Hindi/Tamil/Assamese, themes, text sizing), active gamification counters (2-day streak, 134 gems, Level 4 / 983 XP, Rank #1), patient metadata badge, and daily schedule reminders (hydration, memory session, afternoon medication)."),
        
        ("Figure 2: Cognitive Metrics & 30-Day Trend Analytics Dashboard",
         "Displays quantitative cognitive domain index scores (Memory Index: 88, Focus & Attention: 89, Recall Speed: 80, Reaction Time: 71) alongside dynamic status indicators ('Improving', 'Focus Needed', 'Consistent'). A 30-day SVG line graph tracks score trajectories across all four cognitive categories to evaluate patient improvement."),
        
        ("Figure 3: Interactive 3D Mind Matrix Sanctum Engine (Three.js WebGL)",
         "Features an interactive WebGL 3D cognitive stimulation scene powered by Three.js. Supports orbit rotation, zoom, VR headset immersion mode, particle systems, orbiting target spheres, and real-time inventory tracking (6 Cyber Crystals, 10 Neural Cores, 3 Quantum Rings) with neural matrix charging feedback."),
        
        ("Figure 4: Family Memory Recognition & Reminiscence Exercise Module",
         "Presents photo-based family recognition prompts ('Who is this family member?') with portrait imagery, hints ('Daughter • Loves drinking afternoon tea together...'), and multi-choice selections (Meena, Rahul, Ankit). Detailed profile cards show relationship tags, visiting schedules, and doctor contacts to reinforce familiar identities."),
        
        ("Figure 5: MemoMate Cognitive Gaming Hub Suite",
         "Showcases the 9-game cognitive suite including Quantum Speed Reflex, 3D Dual-N-Back & Number Recall, Spatial Maze Navigator, Acoustic Rhythm & Tone Recall, Card Memory Matrix, Focus Reflex & Math Matrix, Speed Reflex Reaction Test, 3D Holographic Matrix, and 3D Categorical Word Recall with procedural scaling tags and direct play controls.")
    ]

    for idx, (cap, desc) in enumerate(captions):
        img_path = media_files[idx] if idx < len(media_files) else None
        if img_path and os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(2)
            p_img.paragraph_format.keep_with_next = True
            r_i = p_img.add_run()
            r_i.add_picture(img_path, width=Inches(6.8))

            p_c = doc.add_paragraph()
            p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_c.paragraph_format.space_after = Pt(4)
            p_c.paragraph_format.keep_with_next = True
            r_c = p_c.add_run(cap)
            r_c.font.name = 'Times New Roman'
            r_c.font.size = Pt(11)
            r_c.font.bold = True
            r_c.font.italic = True
            r_c.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

            p_d = doc.add_paragraph()
            p_d.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_d.paragraph_format.space_after = Pt(12)
            r_d = p_d.add_run("Description: " + desc)
            r_d.font.name = 'Times New Roman'
            r_d.font.size = Pt(10.5)

    doc.save(docx_path)
    print(f"SUCCESS: Generated Word template -> {docx_path}")

def build_pdf_document(pdf_path, media_files):
    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            self.setStrokeColor(colors.HexColor('#002060'))
            self.setLineWidth(1.5)
            # Outer double border matching template
            self.rect(25, 25, A4[0] - 50, A4[1] - 50)
            self.setLineWidth(0.5)
            self.rect(28, 28, A4[0] - 56, A4[1] - 56)

            # Footer text
            self.setFont("Times-Italic", 9)
            self.setFillColor(colors.HexColor('#333333'))
            self.drawString(36, 34, "SIH 2026 @ Kongu Engineering College | PS ID: SIH26003 | Team MemoMate_KEC")
            self.drawRightString(A4[0] - 36, 34, f"Page {self._pageNumber} of {page_count}")
            self.restoreState()

    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=45
    )

    styles = getSampleStyleSheet()
    
    style_title = ParagraphStyle(
        'HeaderTitle',
        fontName='Times-Bold',
        fontSize=13,
        leading=16,
        alignment=1, # Center
        textColor=colors.HexColor('#002060'),
        spaceAfter=2
    )

    style_sub = ParagraphStyle(
        'HeaderSub',
        fontName='Times-Bold',
        fontSize=11,
        leading=14,
        alignment=1,
        textColor=colors.HexColor('#111111'),
        spaceAfter=8
    )

    style_sec = ParagraphStyle(
        'SecTitle',
        fontName='Times-Bold',
        fontSize=12,
        leading=15,
        alignment=1,
        textColor=colors.HexColor('#002060'),
        spaceBefore=8,
        spaceAfter=8
    )

    style_cap = ParagraphStyle(
        'CapText',
        fontName='Times-BoldItalic',
        fontSize=10.5,
        leading=13,
        alignment=1,
        textColor=colors.HexColor('#002060'),
        spaceBefore=4,
        spaceAfter=3
    )

    style_desc = ParagraphStyle(
        'DescText',
        fontName='Times-Roman',
        fontSize=10,
        leading=13,
        alignment=4, # Justified
        textColor=colors.HexColor('#222222'),
        spaceAfter=10
    )

    story = []

    # Banner Title
    story.append(Paragraph("SMART INDIA HACKATHON 2026 @ KONGU ENGINEERING COLLEGE", style_title))
    story.append(Paragraph("(7th & 8th September 2026)", style_sub))

    # Top Grid Table
    t_left = (
        "<b>Department:</b> Computer Science and Engineering<br/>"
        "<b>Project category:</b> Software<br/>"
        "<b>PS ID:</b> <font color='#002060'><b>SIH26003</b></font><br/>"
        "<b>Team ID:</b> MemoMate_KEC"
    )

    t_right = (
        "<b>Team leader:</b> Srimathi D (Roll No: 22CSR180)<br/>"
        "<b>Team members:</b><br/>"
        "&nbsp;&nbsp;1. Srimathi D (Roll No: 22CSR180)<br/>"
        "&nbsp;&nbsp;2. Team Member 2 (Roll No: 22CSR181)<br/>"
        "&nbsp;&nbsp;3. Team Member 3 (Roll No: 22CSR182)<br/>"
        "&nbsp;&nbsp;4. Team Member 4 (Roll No: 22CSR183)<br/>"
        "&nbsp;&nbsp;5. Team Member 5 (Roll No: 22CSR184)<br/>"
        "<b>Mentor/Co-mentors:</b><br/>"
        "&nbsp;&nbsp;Dr. R. Malathi; Associate Professor, Dept of CSE, KEC<br/>"
        "&nbsp;&nbsp;Prof. P. Suresh; Assistant Professor, Dept of CSE, KEC"
    )

    p_tl = Paragraph(t_left, ParagraphStyle('TL', fontName='Times-Roman', fontSize=10, leading=13))
    p_tr = Paragraph(t_right, ParagraphStyle('TR', fontName='Times-Roman', fontSize=9.5, leading=12))

    tbl = Table([[p_tl, p_tr]], colWidths=[250, 270])
    tbl.setStyle(TableStyle([
        ('BOX', (0, 0), (-1, -1), 1.2, colors.HexColor('#002060')),
        ('INNERGRID', (0, 0), (-1, -1), 0.8, colors.HexColor('#002060')),
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#FAFAFA')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 8))

    story.append(Paragraph("PROTOTYPE SCREENSHOTS & REAL-TIME IMPLEMENTATION", style_sec))

    captions = [
        ("Figure 1: MemoMate Personalized User Dashboard & Daily Schedule Interface",
         "Features personalized greeting ('Good Afternoon, Srimathi'), baseline assessment launch button, AI voice assistant entry point, accessibility toolbar (English/Hindi/Tamil/Assamese, themes, text sizing), active gamification counters (2-day streak, 134 gems, Level 4 / 983 XP, Rank #1), patient metadata badge, and daily schedule reminders."),
        
        ("Figure 2: Cognitive Metrics & 30-Day Trend Analytics Dashboard",
         "Displays quantitative cognitive domain index scores (Memory Index: 88, Focus & Attention: 89, Recall Speed: 80, Reaction Time: 71) alongside dynamic status indicators ('Improving', 'Focus Needed', 'Consistent'). A 30-day SVG line graph tracks score trajectories across all four cognitive categories."),
        
        ("Figure 3: Interactive 3D Mind Matrix Sanctum Engine (Three.js WebGL)",
         "Features an interactive WebGL 3D cognitive stimulation scene powered by Three.js. Supports orbit rotation, zoom, VR headset immersion mode, particle systems, orbiting target spheres, and real-time inventory tracking (6 Cyber Crystals, 10 Neural Cores, 3 Quantum Rings) with neural matrix charging feedback."),
        
        ("Figure 4: Family Memory Recognition & Reminiscence Exercise Module",
         "Presents photo-based family recognition prompts ('Who is this family member?') with portrait imagery, hints ('Daughter • Loves drinking afternoon tea together...'), and multi-choice selections (Meena, Rahul, Ankit). Detailed profile cards show relationship tags, visiting schedules, and doctor contacts."),
        
        ("Figure 5: MemoMate Cognitive Gaming Hub Suite",
         "Showcases the 9-game cognitive suite including Quantum Speed Reflex, 3D Dual-N-Back & Number Recall, Spatial Maze Navigator, Acoustic Rhythm & Tone Recall, Card Memory Matrix, Focus Reflex & Math Matrix, Speed Reflex Reaction Test, 3D Holographic Matrix, and 3D Categorical Word Recall.")
    ]

    for idx, (cap, desc) in enumerate(captions):
        img_path = media_files[idx] if idx < len(media_files) else None
        element_group = []
        if img_path and os.path.exists(img_path):
            img_w = 515
            img_h = 240 if idx in [0, 1, 3] else (300 if idx == 4 else 210)
            element_group.append(Image(img_path, width=img_w, height=img_h))
            element_group.append(Spacer(1, 4))
            element_group.append(Paragraph(cap, style_cap))
            element_group.append(Paragraph("<b>Description:</b> " + desc, style_desc))
            story.append(KeepTogether(element_group))

    doc.build(story, canvasmaker=NumberedCanvas)
    print(f"SUCCESS: Generated PDF template -> {pdf_path}")

def main():
    media_dir = r'C:\Users\LENOVO\.gemini\antigravity-ide\brain\481a6e73-e62a-4a0d-baee-1a7d30e83efb'
    media_files = [
        os.path.join(media_dir, 'media__1788677112588.png'), # Fig 1
        os.path.join(media_dir, 'media__1788677062655.png'), # Fig 2
        os.path.join(media_dir, 'media__1788677112630.png'), # Fig 3
        os.path.join(media_dir, 'media__1788677112655.png'), # Fig 4
        os.path.join(media_dir, 'media__1788677175164.png')  # Fig 5
    ]

    docx_path = r'c:\Users\LENOVO\OneDrive\Desktop\SIH\MemoMate_KEC.docx'
    pdf_path = r'c:\Users\LENOVO\OneDrive\Desktop\SIH\MemoMate_KEC.pdf'

    build_word_document(docx_path, media_files)
    build_pdf_document(pdf_path, media_files)

if __name__ == '__main__':
    main()
