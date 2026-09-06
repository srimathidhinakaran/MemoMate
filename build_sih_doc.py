import os
import glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_page_number_to_run(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def create_document():
    doc = Document()
    
    # ---------------------------------------------------------
    # PAGE SETUP: A4, 1-inch margins on all 4 sides
    # ---------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Header & Footer setup
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("Smart India Hackathon 2026 | MemoMate (PS ID: SIH26003)")
    hrun.font.name = 'Times New Roman'
    hrun.font.size = Pt(10)
    hrun.font.italic = True
    hrun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun1 = fp.add_run("Page ")
    frun1.font.name = 'Times New Roman'
    frun1.font.size = Pt(11)
    frun1.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    add_page_number_to_run(frun1)
    
    # Configure Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(14)
    normal_style.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_style.paragraph_format.space_after = Pt(6)

    # Helpers for headings and text
    def add_p(text, bold_prefix="", space_after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(space_after)
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Times New Roman'
            r_bold.font.size = Pt(14)
            r_bold.font.bold = True
        r_text = p.add_run(text)
        r_text.font.name = 'Times New Roman'
        r_text.font.size = Pt(14)
        return p

    def add_h1(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x00, 0x20, 0x60) # Navy Accent
        return p

    def add_h2(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        return p

    def add_h3(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(title)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        return p

    def add_figure(img_path, caption_text):
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(12)
            p_img.paragraph_format.space_after = Pt(4)
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=Inches(6.2))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(14)
            p_cap.paragraph_format.keep_with_next = True
            r_cap = p_cap.add_run(caption_text)
            r_cap.font.name = 'Times New Roman'
            r_cap.font.size = Pt(12)
            r_cap.font.italic = True
            r_cap.font.bold = True
            r_cap.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Image Paths
    media_dir = r'C:\Users\LENOVO\.gemini\antigravity-ide\brain\481a6e73-e62a-4a0d-baee-1a7d30e83efb'
    img_dashboard = os.path.join(media_dir, 'media__1788677112588.png') # Dashboard
    img_metrics = os.path.join(media_dir, 'media__1788677062655.png')   # Metrics
    img_mind3d = os.path.join(media_dir, 'media__1788677112630.png')    # 3D Sanctum
    img_memories = os.path.join(media_dir, 'media__1788677112655.png')  # My Memories
    img_gamehub = os.path.join(media_dir, 'media__1788677175164.png')   # Game Hub

    # =========================================================
    # COVER PAGE
    # =========================================================
    p_cov_org = doc.add_paragraph()
    p_cov_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_org.paragraph_format.space_before = Pt(36)
    p_cov_org.paragraph_format.space_after = Pt(12)
    r_cov_org = p_cov_org.add_run("SMART INDIA HACKATHON 2026")
    r_cov_org.font.name = 'Times New Roman'
    r_cov_org.font.size = Pt(16)
    r_cov_org.font.bold = True
    r_cov_org.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    p_cov_sub = doc.add_paragraph()
    p_cov_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_sub.paragraph_format.space_after = Pt(36)
    r_cov_sub = p_cov_sub.add_run("Ministry of Education’s Innovation Cell & AICTE")
    r_cov_sub.font.name = 'Times New Roman'
    r_cov_sub.font.size = Pt(13)
    r_cov_sub.font.italic = True

    p_cov_title = doc.add_paragraph()
    p_cov_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_title.paragraph_format.space_after = Pt(18)
    r_cov_title = p_cov_title.add_run("MEMOMATE")
    r_cov_title.font.name = 'Times New Roman'
    r_cov_title.font.size = Pt(28)
    r_cov_title.font.bold = True
    r_cov_title.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    p_cov_subtitle = doc.add_paragraph()
    p_cov_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_subtitle.paragraph_format.space_after = Pt(36)
    r_cov_subtitle = p_cov_subtitle.add_run("AI-POWERED PERSONALIZED COGNITIVE WELLNESS AND MEMORY ASSISTANCE PLATFORM FOR ELDERLY CARE")
    r_cov_subtitle.font.name = 'Times New Roman'
    r_cov_subtitle.font.size = Pt(15)
    r_cov_subtitle.font.bold = True
    r_cov_subtitle.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Details Box Table on Cover Page
    tbl_cov = doc.add_table(rows=9, cols=2)
    tbl_cov.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_cov.autofit = False
    
    cov_data = [
      ("Problem Statement ID:", "SIH26003"),
      ("Problem Statement Title:", "AI-Based Cognitive Gaming and Memory Assistance Platform for Elderly Dementia Patients in North Eastern Region (NER)"),
      ("Project Name / App:", "MemoMate"),
      ("Institution Name:", "Kongu Engineering College"),
      ("Department:", "Department of Computer Science and Engineering"),
      ("Team Member(s):", "Sriman Kumar V (24CSR301) & Srimathi D"),
      ("Faculty Mentor:", "Dr. M. Geetha M.E., Ph.D., Associate Professor, CSE"),
      ("Target Beneficiaries:", "Elderly Individuals, Dementia & MCI Patients, Caregivers, & NER Communities"),
      ("Submission Date:", "September 2026")
    ]
    
    for i, (k, v) in enumerate(cov_data):
        row = tbl_cov.rows[i]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width = Inches(2.2)
        c2.width = Inches(4.0)
        
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(k)
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(12)
        r1.font.bold = True
        
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(v)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
        
        set_cell_background(c1, "F2F4F7")
        set_cell_background(c2, "FFFFFF")
        set_cell_margins(c1, top=80, bottom=80, left=100, right=100)
        set_cell_margins(c2, top=80, bottom=80, left=100, right=100)

    doc.add_page_break()

    # =========================================================
    # 1. ABSTRACT
    # =========================================================
    add_h1("1. ABSTRACT")
    add_p("Cognitive decline, Mild Cognitive Impairment (MCI), and early-stage dementia present severe challenges to the aging population worldwide, particularly in remote and culturally distinct areas such as the North Eastern Region (NER) of India. Traditional cognitive intervention tools often fail due to lack of personalization, complex non-intuitive user interfaces, generic non-adaptive activity sets, and language barriers that alienate non-English speaking elderly users. This paper presents MemoMate, a personalized, AI-driven cognitive gaming and memory assistance platform engineered specifically for elderly dementia care and cognitive maintenance.")
    add_p("MemoMate combines interactive WebGL 3D cognitive stimulation exercises, an adaptive cognitive performance monitoring engine, natural voice assistance ('Talk to MemoMate'), localized North Eastern regional language support (English, Hindi, Tamil, Assamese), and gamified habit formation systems. The system establishes a baseline cognitive profile spanning four primary domains: Visual Memory, Executive Focus & Attention, Sequential Recall Speed, and Motor Reaction Velocity. Based on real-time performance telemetry, MemoMate's AI recommendation engine dynamically tailors activity sets, adjusts game difficulty, tracks weekly metrics, and fosters emotional well-being through an evolving digital Memory Garden and family reminiscence exercises.")
    add_p("Full-stack data persistence powered by a secure Node.js/Express backend and MongoDB Atlas guarantees strict per-user authorization, privacy, and continuous long-term telemetry tracking. MemoMate serves as a non-invasive, accessible, non-pharmacological supportive platform designed to enhance daily engagement, support caregiver monitoring, and preserve cognitive dignity for senior citizens.")

    # =========================================================
    # 2. INTRODUCTION
    # =========================================================
    add_h1("2. INTRODUCTION")
    add_p("As healthcare advancements increase human longevity globally, the prevalence of age-related neurodegenerative conditions—specifically Alzheimer’s disease and various forms of dementia—has risen exponentially. In India, over 5.3 million individuals aged 60 and above live with dementia, a number projected to exceed 14 million by 2050. The North Eastern Region (NER) of India faces unique healthcare delivery hurdles due to geographic isolation, limited specialized geriatric clinical infrastructure, and rich linguistic diversity comprising over 200 dialects.")
    add_p("Cognitive stimulation therapy (CST) and structured mental exercises are clinically proven non-pharmacological interventions that help slow cognitive decline, enhance neuroplasticity, and maintain functional independence. However, conventional digital brain training solutions are designed primarily for younger, tech-savvy demographics. They suffer from high cognitive load UI, small touch targets, rapid timer pressures, lack of regional language options, and monotonous repetitive tasks that cause user frustration and abandonment.")
    add_p("MemoMate addresses these critical bottlenecks by introducing an empathetic, mobile-first digital environment tailored specifically for elderly usability. Designed with high-contrast themes, scalable typography, multi-modal voice interaction, multi-dialect support, and positive reinforcement gamification (XP, levels, daily streaks, virtual garden growth), MemoMate transforms routine cognitive exercises into an engaging daily routine.")

    # =========================================================
    # 3. PROBLEM STATEMENT
    # =========================================================
    add_h1("3. PROBLEM STATEMENT")
    add_p("SIH Problem Statement Title:", "Official Title: ")
    add_p("AI-Based Cognitive Gaming and Memory Assistance Platform for Elderly Dementia Patients in North Eastern Region (NER)", space_after=12)
    add_p("Problem Statement ID:", "Official ID: ")
    add_p("SIH26003", space_after=12)
    
    add_h2("3.1 Background and Problem Analysis")
    add_p("Elderly individuals suffering from early-stage dementia or mild cognitive impairment exhibit progressive deficits in short-term memory, attention span, orientation, spatial navigation, and processing velocity. In the North Eastern Region of India, additional systemic constraints aggravate these challenges:")
    add_p("1. Lack of Localized Tools: Existing digital applications lack support for NER regional languages and local cultural contexts (such as traditional NER festivals, crafts, and heritage imagery), leading to cultural detachment.", bold_prefix="")
    add_p("2. Usability & Accessibility Barriers: Standard smartphones and applications feature complex navigation, small buttons (less than 44px), and intricate gesture requirements that frustrate elderly users with tremors or reduced motor control.", bold_prefix="")
    add_p("3. Static & Generic Content: Conventional brain games apply fixed difficulty levels and generic puzzles, failing to adapt to individual patient performance fluctuations or specific cognitive weaknesses.", bold_prefix="")
    add_p("4. Isolated Caregiver Experience: Caregivers lack real-time visibility into daily cognitive performance trends, compliance with medication/hydration schedules, or early indicators of sharp attention drops.", bold_prefix="")

    # =========================================================
    # 4. EXISTING SYSTEM & COMPARATIVE ANALYSIS
    # =========================================================
    add_h1("4. EXISTING SYSTEM ANALYSIS")
    add_p("Current approaches to cognitive maintenance fall into two categories: traditional offline physical activities (such as paper crosswords, physical flashcards, and manual memory games) and commercial digital brain-training applications. Table 1 provides a comprehensive comparative evaluation highlighting the gaps in existing solutions and how MemoMate bridges them.")

    add_h2("Table 1: Comparative Analysis of Existing Solutions vs. MemoMate")
    tbl_comp = doc.add_table(rows=6, cols=3)
    tbl_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    comp_headers = ["Feature Domain", "Existing Digital Applications", "MemoMate Proposed Platform"]
    for j, h in enumerate(comp_headers):
        cell = tbl_comp.rows[0].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = cell.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "002060")
        set_cell_margins(cell, 100, 100, 120, 120)

    comp_rows = [
        ("Language & Dialect Support", "Predominantly English; zero regional NER dialect localization.", "Native support for English, Hindi, Tamil, and Assamese (as-IN) with full UI translation."),
        ("User Interface Design", "Dense layouts, small text, complex multi-touch gestures, high visual noise.", "Elderly-first responsive design, high contrast, 46px+ touch targets, font-scaler."),
        ("Personalization Engine", "Static difficulty curves; generic puzzle distribution across all users.", "AI-driven baseline profile & adaptive difficulty scaling across 4 cognitive domains."),
        ("Voice Assistance", "Text-only inputs or basic English voice controls.", "Integrated 'Talk to MemoMate' assistant with dynamic SpeechSynthesis & Recognition locale mapping."),
        ("Gamification & Habit Building", "Intrusive paywalls, competitive stress, complex economy.", "Positive habit reinforcement via XP, levels, daily streaks, and digital Memory Garden expansion.")
    ]

    for i, row_data in enumerate(comp_rows):
        row = tbl_comp.rows[i+1]
        bg = "F9FAFB" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            if j == 0:
                cell.width = Inches(1.8)
            elif j == 1:
                cell.width = Inches(2.2)
            else:
                cell.width = Inches(2.2)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = cell.paragraphs[0].add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
            set_cell_background(cell, bg)
            set_cell_margins(cell, 80, 80, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # =========================================================
    # 5. PROPOSED SYSTEM — MEMOMATE
    # =========================================================
    add_h1("5. PROPOSED SYSTEM — MEMOMATE")
    add_p("MemoMate is a holistic, multi-modal software ecosystem designed to provide accessible, non-pharmacological cognitive support. At its core, MemoMate operationalizes an interactive closed-loop architecture where user performance continuously informs system behavior:")
    add_p("User Onboarding & Authentication → Baseline Cognitive Assessment → Real-Time Telemetry & Domain Scoring → AI Recommendation Engine → Tailored Cognitive Game Selection → Dynamic Difficulty Adjustment → Multi-Modal Voice & Reminiscence Interaction → Gamified Reward & Memory Garden Growth → Caregiver Analytics Monitoring.", bold_prefix="System Conceptual Architecture: ", space_after=12)
    add_p("By combining high-performance WebGL 3D graphics (Three.js), speech recognition pipelines, regional language dictionaries, and a database-driven REST API, MemoMate delivers a dignified and therapeutic user experience.")

    # Insert Figure 1
    add_figure(img_dashboard, "Figure 1: MemoMate Personalized User Dashboard & Daily Schedule Interface")
    add_p("Figure 1 demonstrates the MemoMate primary user dashboard. The interface features a prominent personalized greeting ('Good Afternoon, Srimathi'), direct action buttons for 'Start Baseline Assessment' and 'Talk to MemoMate', a high-visibility accessibility bar, active gamification counters (2-day streak, 134 gems/points, Level 4 / 983 XP, Rank #1), patient profile metadata, and an interactive schedule tracking daily hydration and cognitive activity completion status.")

    # =========================================================
    # 6. OBJECTIVES
    # =========================================================
    add_h1("6. SYSTEM OBJECTIVES")
    add_p("The primary technical and operational objectives of the MemoMate project include:")
    objectives = [
        "To establish a comprehensive, non-invasive digital baseline cognitive assessment evaluating Memory, Attention, Recall, and Reaction Time.",
        "To engineer an adaptive AI-driven recommendation engine that tailors game selection and difficulty curves based on individual cognitive performance indicators.",
        "To provide 100% localized multilingual interface consistency supporting English, Hindi, Tamil, and Assamese (as-IN) to dismantle linguistic barriers in the NER region.",
        "To integrate a real-time speech synthesis and speech recognition assistant ('Talk to MemoMate') capable of localized conversational interaction.",
        "To foster sustained daily engagement among elderly patients through positive reinforcement gamification including XP points, leveling, streaks, and digital Memory Garden growth.",
        "To implement a specialized Reminiscence and Family Memory Recognition module to reinforce familiarity with close relatives and personal identity.",
        "To deliver a mobile-first, elderly-friendly user interface incorporating large touch targets (>=46px), high contrast themes, and dynamic font scaling.",
        "To provide secure, isolated RESTful backend data persistence utilizing Node.js, Express, and MongoDB Atlas with encrypted password hashing.",
        "To offer caregivers actionable real-time cognitive trend analytics and automated alerts regarding drastic performance dips or missed routine schedules."
    ]
    for idx, obj in enumerate(objectives, 1):
        add_p(obj, bold_prefix=f"{idx}. ", space_after=4)

    # =========================================================
    # 7. KEY FEATURES & MODULE ANALYSIS
    # =========================================================
    add_h1("7. KEY FEATURES & IMPLEMENTED MODULES")
    
    add_h2("7.1 User Registration and Real MongoDB Authentication")
    add_p("MemoMate implements a secure, database-backed authentication system. User registration captures essential demographic metadata (Name, Age, Role, Preferred Language, Preferred Theme) and enforces strict validation (minimum 6-character passwords, unique email constraints). Passwords are securely hashed using bcrypt before storage in MongoDB Atlas, ensuring zero plain-text data exposure.")

    add_h2("7.2 Baseline Cognitive Assessment & Performance Telemetry")
    add_p("Upon initial login, elderly users are guided through an unhurried baseline assessment evaluating four primary cognitive domains:")
    add_p("• Visual & Spatial Memory: Assessed via pattern recall, matrix matching, and WebGL 3D cube alignment.", space_after=2)
    add_p("• Executive Focus & Attention: Evaluated through target tracking, orb detection, and selective focus tasks.", space_after=2)
    add_p("• Sequential Recall Speed: Evaluated using multi-digit N-back digit sequences and categorical word recall.", space_after=2)
    add_p("• Motor Reaction Velocity: Measured in milliseconds during WebGL 3D orb tap reflexes and visual stimulus tests.", space_after=6)

    # Insert Figure 2
    add_figure(img_metrics, "Figure 2: Cognitive Profile Metrics and 30-Day Trend Analytics Dashboard")
    add_p("Figure 2 illustrates the MemoMate Metrics & Trends module. The top panel displays quantitative domain scores (Memory Index: 88, Focus & Attention: 89, Recall Speed: 80, Reaction Time: 71) alongside dynamic trend badges ('Improving', 'Focus Needed', 'Consistent'). The central SVG graph plots multi-week cognitive performance trajectories across all four domains, giving patients and caregivers clear visual feedback.")

    add_h2("7.3 Interactive WebGL 3D Mind Matrix Sanctum")
    add_p("To provide rich sensory stimulation, MemoMate incorporates a custom Three.js WebGL 3D engine ('Mind Matrix Sanctum'). Users interact with orbiting 3D celestial bodies, illuminated neural cores, and quantum rings. The module supports real-time mouse/touch drag orbiting, smooth camera zooming, and an optional Virtual Reality (VR) headset mode for immersive cognitive therapy.")

    # Insert Figure 3
    add_figure(img_mind3d, "Figure 3: Interactive 3D Mind Matrix Sanctum Engine (Three.js WebGL Rendering)")
    add_p("Figure 3 showcases the interactive 3D Mind Matrix Sanctum in execution. Rendered using Three.js WebGL shaders, the sanctum displays orbiting target spheres, real-time particle effects, VR toggle controls, and resource counters (6 Cyber Crystals, 10 Neural Cores, 3 Quantum Rings) while executing neural matrix charging animations.")

    add_h2("7.4 Reminiscence & Family Memory Recognition")
    add_p("Memory loss in dementia patients heavily affects family recognition. MemoMate includes a dedicated 'My Memories' reminiscence module that presents photo-based identification prompts featuring registered family members (e.g., Daughter, Son, Grandson, Family Doctor) along with personalized relationship hints and visit schedules.")

    # Insert Figure 4
    add_figure(img_memories, "Figure 4: Family Memory Recognition & Reminiscence Exercise Module")
    add_p("Figure 4 displays the Family Memory Recognition module. The top card presents an active identification prompt ('Who is this family member?') with portrait photo and contextual hints ('Daughter • Loves drinking afternoon tea together...'), flanked by multiple-choice selection buttons (Meena, Rahul, Ankit). Below, detailed cards display family profile entries, schedules, and clinical doctor contacts.")

    add_h2("7.5 Comprehensive Cognitive Gaming Hub")
    add_p("MemoMate features a diverse suite of 9 procedural and dynamic 3D cognitive games targeting distinct neurological pathways. Games utilize adaptive scoring algorithms that dynamically scale target counts, sequence lengths, and reaction timeouts based on real-time accuracy.")

    # Insert Figure 5
    add_figure(img_gamehub, "Figure 5: MemoMate Cognitive Gaming Hub Suite")
    add_p("Figure 5 displays the Cognitive Exercises hub. The grid showcases 9 specialized cognitive games: Quantum Speed Reflex, 3D Dual-N-Back & Number Recall, Spatial Maze Navigator, Acoustic Rhythm & Tone Recall, Card Memory Matrix, Focus Reflex & Math Matrix, Speed Reflex Reaction Test, 3D Holographic Matrix, and 3D Categorical Word Recall. Each game card highlights category tags, procedural execution badges, and direct launch controls.")

    add_h2("7.6 Digital Memory Garden & Habit Gamification")
    add_p("To foster long-term compliance without stress, MemoMate introduces a digital Memory Garden. Completing cognitive sessions grants watering points that blossom virtual flowers, plants, and trees. Streak freezes protect daily progress during rest days, while leveling up unlocks customizable mind matrix relics.")

    add_h2("7.7 'Talk to MemoMate' Multilingual AI Voice Assistant")
    add_p("The platform features an integrated voice assistant ('Talk to MemoMate') accessible from every view. Built on Web Speech API standard locales, the assistant dynamically maps speech recognition (`SpeechRecognition.lang`) and text-to-speech engines (`SpeechSynthesis.lang`) to match the user’s selected language (English: `en-US`, Hindi: `hi-IN`, Tamil: `ta-IN`, Assamese: `as-IN`). If a device lacks native synthesis voices for a regional dialect, MemoMate gracefully notifies the user in their selected language rather than silently reverting to English.")

    # =========================================================
    # 8. SYSTEM WORKFLOW & ARCHITECTURE
    # =========================================================
    add_h1("8. SYSTEM WORKFLOW AND TECHNICAL ARCHITECTURE")
    add_p("The MemoMate system architecture follows a decoupled, client-server model engineered for responsiveness, cross-platform compatibility, and secure data handling.")

    add_h2("8.1 End-to-End Data Flow")
    add_p("1. Client Layer: The user accesses MemoMate via a mobile browser or desktop PWA wrapper. User actions (taps, voice commands, game inputs) are captured by React components.", bold_prefix="")
    add_p("2. Service & Context Layer: `AuthContext` and `api.js` manage global application state, authentication tokens (JWT), active language translations (`NER_TRANSLATIONS`), and theme styling.", bold_prefix="")
    add_p("3. API Gateway / Router Layer: HTTP requests are transmitted to the Express REST API (`backend/server.js` or Vercel serverless `/api/*` endpoints) using Axios with Bearer token headers.", bold_prefix="")
    add_p("4. Controller & Business Logic Layer: Express controllers (`authController`, `sessionController`, `cognitiveController`) validate request payloads, calculate exponentially weighted cognitive scores, update streak multipliers, and invoke recommendation algorithms.", bold_prefix="")
    add_p("5. Persistence Layer: Mongoose ORM interacts with MongoDB Atlas to execute atomic CRUD operations across structured collections (`users`, `cognitiveprofiles`, `cognitivesessions`, `gamifications`, `gardenprogresses`).", bold_prefix="")

    add_h2("8.2 Technology Stack")
    add_p("MemoMate is built strictly using modern open-source web technologies:")
    
    tbl_tech = doc.add_table(rows=6, cols=3)
    tbl_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    t_headers = ["Layer / Domain", "Technology Choice", "Functional Role in MemoMate"]
    for j, h in enumerate(t_headers):
        cell = tbl_tech.rows[0].cells[j]
        r = cell.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "002060")
        set_cell_margins(cell, 100, 100, 120, 120)

    t_data = [
        ("Frontend UI Framework", "React 18 + Vite", "High-performance SPA rendering, reactive state management, fast HMR build pipeline."),
        ("3D Engine & VR", "Three.js (WebGL)", "Real-time 3D Mind Matrix Sanctum rendering, particle systems, and VR camera controls."),
        ("Styling & Responsive UI", "Vanilla CSS Design Tokens", "Custom dark themes, high-contrast mode, 46px+ touch target CSS rules, font scalers."),
        ("Backend Server API", "Node.js + Express.js", "RESTful API endpoints, JWT authentication middleware, Vercel Serverless Function handler."),
        ("Database Storage", "MongoDB Atlas + Mongoose", "Cloud NoSQL document database, schema modeling, indexed user data isolation.")
    ]

    for i, row_data in enumerate(t_data):
        row = tbl_tech.rows[i+1]
        bg = "F9FAFB" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = cell.paragraphs[0].add_run(val)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
            set_cell_background(cell, bg)
            set_cell_margins(cell, 80, 80, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # =========================================================
    # 9. DATABASE DESIGN
    # =========================================================
    add_h1("9. DATABASE DESIGN & SCHEMA ARCHITECTURE")
    add_p("Data persistence in MemoMate is modeled around MongoDB Atlas collections connected via Mongoose schemas. User data isolation is enforced by embedding or referencing `userId` (MongoDB `ObjectId`) across all sub-collections.")

    add_h2("9.1 Core Database Collections")
    add_p("• `users`: Stores core account credentials (`email`, `password` hashed with bcrypt), demographic data (`name`, `age`, `role`), and UI preferences (`preferredLanguage`, `preferredTheme`).", space_after=2)
    add_p("• `cognitiveprofiles`: Tracks cumulative domain scores (`memoryScore`, `attentionScore`, `recallScore`, `reactionScore`, `overallScore`) and assessment completion flags (`assessed: true/false`).", space_after=2)
    add_p("• `cognitivesessions`: Logs detailed individual game session telemetry including `activity` name, `category`, `difficulty`, `score` (0–100), `accuracy` percentage, `reactionTime` (ms), and ISO timestamp.", space_after=2)
    add_p("• `gamifications`: Maintains XP points, gem currency, level status, current streak, highest streak, last active check-in date, unlocked badges, and daily quest progression.", space_after=2)
    add_p("• `gardenprogresses`: Stores digital garden counters (`plants`, `flowers`, `trees`, `streak`, `totalActivities`) linked to user account IDs.", space_after=6)

    # =========================================================
    # 10. ADAPTIVE PERSONALIZATION & AI RECOMMENDATION
    # =========================================================
    add_h1("10. ADAPTIVE PERSONALIZATION ENGINE")
    add_p("MemoMate avoids fixed, repetitive game schedules by implementing an exponentially weighted score update model and a dynamic weak-area identification algorithm.")

    add_h2("10.1 Score Smoothing Algorithm")
    add_p("When a user completes a cognitive activity, their domain score is updated using an exponentially weighted moving average (EWMA) formula:")
    add_p("New_Domain_Score = (Previous_Domain_Score × 0.65) + (Current_Session_Score × 0.35)", bold_prefix="Formula: ", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_p("This formula ensures that single anomalous poor performances (due to temporary fatigue) do not drastically ruin a user's metric, while consistent gains over 3–5 sessions steadily reflect genuine cognitive improvement.")

    add_h2("10.2 Recommendation Logic")
    add_p("The recommendation engine compares current domain scores (`memoryScore`, `attentionScore`, `recallScore`, `reactionScore`). The domain exhibiting the lowest score is automatically flagged as the primary target area. The engine then selects an appropriate game activity and sets initial difficulty ('Easy' for scores < 60, 'Medium' for 60–85, 'Hard' for > 85), prompting the user directly on their dashboard hero card.")

    # =========================================================
    # 11. MULTILINGUAL & VOICE PIPELINE
    # =========================================================
    add_h1("11. MULTILINGUAL & VOICE PIPELINE")
    add_p("To eliminate linguistic exclusion in the North Eastern Region of India, MemoMate implements a centralized internationalization system (`NER_TRANSLATIONS`) supporting English, Hindi, Tamil, and Assamese (as-IN).")

    add_h2("11.1 Zero English Leakage Consistency")
    add_p("All user interface elements—including page titles, navigation labels, game instructions, HUD counters, modal alerts, dynamic feedback toasts, and voice assistant responses—are driven by centralized key lookups (`t('key', { params })`). String interpolation replaces dynamic values (e.g. score counts, level numbers) within localized template strings, ensuring 100% language purity without accidental English fallbacks.")

    add_h2("11.2 Speech Recognition & TTS Locale Mapping")
    add_p("The 'Talk to MemoMate' assistant leverages native Web Speech APIs with explicit locale mapping:")
    add_p("• English: `en-US` | Hindi: `hi-IN` | Tamil: `ta-IN` | Assamese: `as-IN`", bold_prefix="Locale Mapping: ", space_after=6)
    add_p("When speech recognition activates, `recognition.lang` is set to match the active application language. Similarly, `SpeechSynthesisUtterance.lang` filters installed device voices to match the localePrefix. If an exact voice is unavailable on a specific operating system, MemoMate informs the user in their active language rather than outputting mispronounced English speech.")

    # =========================================================
    # 12. RESPONSIVE MOBILE DESIGN & ACCESSIBILITY
    # =========================================================
    add_h1("12. ELDERLY ACCESSIBILITY AND RESPONSIVE DESIGN")
    add_p("Designing software for elderly dementia patients requires strict adherence to specialized UI/UX heuristics:")
    add_p("1. Touch Target Calibration: All interactive buttons, card selection areas, and game controls feature a minimum height and width of 46px (exceeding standard W3C 44px recommendations) to accommodate age-related tremors or reduced dexterity.", bold_prefix="")
    add_p("2. High Contrast Theme Engine: Includes 4 built-in visual modes ('Calm Nature', 'Healthcare Dark', 'Fire Pro High Contrast', 'Daylight Clean') to alleviate visual fatigue and support varying visual acuities.", bold_prefix="")
    add_p("3. Dynamic Font Scaling: Users can toggle text size between 'Normal' (14pt body), 'Large' (16pt body), and 'Extra Large' (18pt body) instantly across all routes.", bold_prefix="")
    add_p("4. Script Expansion Safety: Non-Latin scripts (Tamil, Hindi, Assamese) typically occupy greater horizontal width. MemoMate utilizes `word-break: break-word` and flexible flexbox/grid containers to prevent layout clipping or horizontal scrollbars on 360px–414px mobile viewports.", bold_prefix="")

    # =========================================================
    # 13. SECURITY, PRIVACY & SCALABILITY
    # =========================================================
    add_h1("13. SECURITY, PRIVACY AND SCALABILITY")
    add_p("MemoMate prioritizes patient data security and architectural scalability:")
    add_p("• Password Security: User passwords are encrypted using bcrypt hashing (salt factor 10). Plain-text passwords are never stored or logged.", space_after=2)
    add_p("• Authorization & JWT: API routes enforce JSON Web Token (JWT) verification via HTTP Bearer headers, ensuring users can only read or write their own cognitive records.", space_after=2)
    add_p("• Zero Secret Exposure: Environment credentials (`MONGO_URI`, `JWT_SECRET`, `GROQ_API_KEY`) are managed strictly via `.env` files and Vercel serverless environment variables, ensuring zero secret leakage in client bundles.", space_after=2)
    add_p("• Stateless Serverless Scalability: Deployed on Vercel Serverless Functions (`api/index.js`), the backend scales automatically to handle concurrent user traffic with cached Mongoose database connections.", space_after=6)

    # =========================================================
    # 14. UNIQUENESS & INNOVATION (SIH EVALUATION)
    # =========================================================
    add_h1("14. UNIQUENESS AND INNOVATION")
    add_p("MemoMate stands out in the SIH 2026 evaluation through several key innovations:")
    add_p("1. NER Regional Language Focus: Native Assamese (as-IN) integration alongside Hindi, Tamil, and English dismantles linguistic isolation for elderly populations in North Eastern India.", bold_prefix="")
    add_p("2. Integrated WebGL 3D & VR Stimulation: Leverages WebGL Three.js spatial scenes for immersive 3D visual and spatial memory therapy.", bold_prefix="")
    add_p("3. Multi-Modal Interaction Pipeline: Combines direct touch, dynamic voice assistance ('Talk to MemoMate'), visual reminiscence, and dynamic audio-rhythm games in a unified platform.", bold_prefix="")
    add_p("4. Non-Stress Habit Formation: Uses digital Memory Garden growth, streak freezes, and daily XP rewards to motivate cognitive exercise compliance without creating anxiety.", bold_prefix="")
    add_p("5. Holistic Caregiver Telemetry: Provides real-time caregiver trend dashboards tracking multi-week performance, hydration compliance, and attention alerts.", bold_prefix="")

    # =========================================================
    # 15. ADVANTAGES AND LIMITATIONS
    # =========================================================
    add_h1("15. ADVANTAGES AND LIMITATIONS")
    
    add_h2("15.1 Platform Advantages")
    add_p("• Non-invasive, non-pharmacological supportive tool for elderly cognitive maintenance.", space_after=2)
    add_p("• 100% localized multilingual consistency eliminating English confusion for elderly users.", space_after=2)
    add_p("• Fully mobile-responsive layout optimized for smartphones, tablets, and desktop devices.", space_after=2)
    add_p("• Secure cloud persistence allowing patients and caregivers to access history seamlessly.", space_after=6)

    add_h2("15.2 System Limitations")
    add_p("• Voice synthesis quality depends on native operating system / device TTS engine installations.", space_after=2)
    add_p("• Requires internet connectivity for MongoDB Atlas cloud synchronization and speech recognition APIs.", space_after=2)
    add_p("• MemoMate provides supportive cognitive wellness tracking and is NOT a certified clinical diagnostic tool.", space_after=6)

    # =========================================================
    # 16. FUTURE ENHANCEMENTS
    # =========================================================
    add_h1("16. FUTURE ENHANCEMENTS")
    add_p("Planned future roadmap enhancements for MemoMate include:")
    add_p("1. Expanded NER Dialect Support: Adding native translation dictionaries for additional North Eastern regional languages including Manipuri (Meitei), Bodo, Khasi, and Mizu.", bold_prefix="")
    add_p("2. Offline PWA Synchronization: Integrating IndexedDB local storage and Web Background Sync to enable complete offline gameplay in rural areas with intermittent internet connectivity.", bold_prefix="")
    add_p("3. EEG & Wearable Integration: Exploring integration with consumer EEG headbands (e.g. Muse) to record real-time brainwave attention/meditation telemetry during 3D exercises.", bold_prefix="")
    add_p("4. Clinician Export Portal: Generating automated PDF summary reports formatted for geriatric neurologists and clinical caregivers.", bold_prefix="")

    # =========================================================
    # 17. SUSTAINABLE DEVELOPMENT GOALS (SDGs) ALIGNMENT
    # =========================================================
    add_h1("17. ALIGNMENT WITH SUSTAINABLE DEVELOPMENT GOALS")
    add_p("MemoMate directly supports key United Nations Sustainable Development Goals (SDGs):")
    add_p("• SDG 3: Good Health and Well-Being — Promotes healthy aging and mental well-being for elderly citizens suffering from cognitive decline.", bold_prefix="")
    add_p("• SDG 10: Reduced Inequalities — Dismantles technological and linguistic barriers by providing free, accessible regional language cognitive tools for underserved communities in North Eastern India.", bold_prefix="")
    add_p("• SDG 9: Industry, Innovation, and Infrastructure — Demonstrates novel application of WebGL 3D graphics, speech recognition, and AI personalization in healthcare technology.", bold_prefix="")

    # =========================================================
    # 18. CONCLUSION
    # =========================================================
    add_h1("18. CONCLUSION")
    add_p("MemoMate represents a comprehensive, empathetic, and innovative technological response to SIH Problem Statement SIH26003. By integrating adaptive cognitive exercises, WebGL 3D sensory stimulation, localized regional language support (English, Hindi, Tamil, Assamese), voice interaction, and positive habit gamification into a unified mobile-friendly platform, MemoMate provides elderly individuals with a dignified tool to exercise their mental faculties daily. Backed by a robust Node.js/Express backend and MongoDB Atlas persistence, MemoMate bridges the gap between digital healthcare innovation and senior citizen accessibility.")

    # =========================================================
    # 19. REFERENCES
    # =========================================================
    add_h1("19. REFERENCES")
    refs = [
        "Ministry of Education’s Innovation Cell & AICTE, Smart India Hackathon (SIH 2026) Problem Statement SIH26003 Documentation.",
        "World Health Organization (WHO), 'Global Action Plan on the Public Health Response to Dementia 2017–2025', Geneva, 2017.",
        "Dementia India Report, Alzheimer’s and Related Disorders Society of India (ARDSI), New Delhi, 2020.",
        "React JS Documentation & Hooks API Reference, Meta Open Source, 2024. Available: https://react.dev",
        "Three.js WebGL 3D Computer Graphics Library Documentation, Ricardo Cabello (Mr.doob), 2024. Available: https://threejs.org",
        "MongoDB Atlas & Mongoose ODM Documentation, MongoDB Inc., 2024. Available: https://www.mongodb.com/docs",
        "W3C Web Accessibility Initiative (WAI), 'Mobile Accessibility: User Experiences and Behaviors Guidelines', 2023."
    ]
    for idx, ref in enumerate(refs, 1):
        add_p(ref, bold_prefix=f"[{idx}] ", space_after=4)

    # Save Document
    out_path = r'c:\Users\LENOVO\OneDrive\Desktop\SIH\MEMOMATE_SIH_2026_DOCUMENTATION.docx'
    doc.save(out_path)
    print(f"SUCCESS: Created {out_path}")

if __name__ == '__main__':
    create_document()
